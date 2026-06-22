#!/usr/bin/env python3
"""Export an agi-qa series post to the Word review format used on the Desktop.

Usage: python3 export-review-docx.py <post.md> [--final]

Formatting matches "Richard Sutton on Experience.docx":
- Title style for the post title
- gray (666666) italic 12pt preamble (description, researchers, sources)
- navy (1F4E79) bold "Q1  " prefix + black bold 12pt question lines
- plain 11pt body, bold thinker headers, inline bold/italic preserved
- 9pt gray italic source notes with bold "Source: " prefix
- navy italic paraphrase blocks with bold "Paraphrase: " prefix
Straight quotes are converted to curly quotes (markdown source keeps straight).
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

GRAY = RGBColor(0x66, 0x66, 0x66)
NAVY = RGBColor(0x1F, 0x4E, 0x79)


def smarten(text):
    text = re.sub(r"(\w)'(\w)", "’".join((r"\1", r"\2")), text)
    text = re.sub(r'(?<=\S)"', "”", text)  # close after non-space first
    text = text.replace('"', "“")
    text = re.sub(r"(?<=\S)'", "’", text)
    text = text.replace("'", "‘")
    return text


def inline_runs(text):
    """Split markdown text into (text, bold, italic) runs."""
    out = []
    for piece in re.split(r"(\*\*.*?\*\*|\*.*?\*)", text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            out.append((piece[2:-2], True, False))
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            out.append((piece[1:-1], False, True))
        else:
            out.append((piece, False, False))
    return out


def add_runs(p, pieces, size=None, color=None, force_italic=False):
    for text, bold, italic in pieces:
        r = p.add_run(smarten(text))
        r.bold = bold or None
        r.italic = (italic or force_italic) or None
        if size:
            r.font.size = size
        if color:
            r.font.color.rgb = color


def style_spacing(p):
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15


def main():
    src = Path(sys.argv[1])
    raw = src.read_text()
    m = re.match(r"---\n(.*?)\n---\n", raw, re.S)
    front, body = m.group(1), raw[m.end():]
    title = re.search(r"title:\s*'(.*)'", front).group(1)

    doc = Document()
    tp = doc.add_paragraph(style="Title")
    tp.add_run(smarten(title))

    seen_q = False
    for block in re.split(r"\n\n+", body.strip()):
        raw_block = block.strip()
        if not raw_block:
            continue
        if raw_block.startswith("|"):
            lines = [ln for ln in raw_block.splitlines()
                     if not re.match(r"^\|[\s:|-]+\|$", ln.strip())]
            rows = [[c.strip() for c in ln.strip().strip("|").split("|")]
                    for ln in lines]
            ncols = len(rows[0])
            rows = [r[:ncols] + [""] * (ncols - len(r)) for r in rows]
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    cp = table.rows[ri].cells[ci].paragraphs[0]
                    add_runs(cp, inline_runs(cell), size=Pt(10))
                    if ri == 0:
                        for r in cp.runs:
                            r.bold = True
                            r.font.color.rgb = NAVY
            doc.add_paragraph()
            continue
        if raw_block.startswith("<figure"):
            cap = re.search(r"<figcaption>(.*?)</figcaption>", raw_block, re.S)
            caption = " ".join(cap.group(1).split()) if cap else "(figure)"
            fm = re.match(r"(Figure \d+)\.\s*(.*)", caption, re.S)
            label, rest = (fm.group(1), fm.group(2)) if fm else ("Figure", caption)
            tm = re.search(r'<p class="chart-title">(.*?)</p>', raw_block, re.S)
            chart_title = f" — {tm.group(1).strip()}" if tm else ""
            p = doc.add_paragraph()
            style_spacing(p)
            r = p.add_run(f"{label}{chart_title} (chart appears here on the blog)")
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = NAVY
            p2 = doc.add_paragraph()
            style_spacing(p2)
            add_runs(p2, inline_runs(rest), size=Pt(9), color=GRAY,
                     force_italic=True)
            continue
        if raw_block.startswith("<details"):
            sm = re.search(r"<summary>(.*?)</summary>", raw_block, re.S)
            p = doc.add_paragraph()
            style_spacing(p)
            r = p.add_run("▸ " + (sm.group(1).strip() if sm else "Side quest"))
            r.bold = True
            r.font.color.rgb = NAVY
            inner = re.sub(r"<summary>.*?</summary>", "", raw_block, flags=re.S)
            for tag, content in re.findall(r"<(p|table)[^>]*>(.*?)</\1>", inner, re.S):
                if tag == "p":
                    pp = doc.add_paragraph()
                    style_spacing(pp)
                    pp.paragraph_format.left_indent = Pt(18)
                    em = re.match(r"\s*<em>(.*)</em>\s*$", content, re.S)
                    add_runs(pp, inline_runs(em.group(1) if em else content),
                             force_italic=bool(em))
                else:
                    rows = [re.findall(r"<td>(.*?)</td>", tr, re.S)
                            for tr in re.findall(r"<tr>(.*?)</tr>", content, re.S)]
                    ncols = max(len(r) for r in rows)
                    rows = [r + [""] * (ncols - len(r)) for r in rows]
                    table = doc.add_table(rows=len(rows), cols=ncols)
                    table.style = "Table Grid"
                    for ri, row in enumerate(rows):
                        for ci, cell in enumerate(row):
                            cp = table.rows[ri].cells[ci].paragraphs[0]
                            add_runs(cp, inline_runs(cell.strip()), size=Pt(10))
                            if ri == 0:
                                for rr in cp.runs:
                                    rr.bold = True
                                    rr.font.color.rgb = NAVY
                    doc.add_paragraph()
            continue
        if raw_block.startswith("> "):
            p = doc.add_paragraph()
            style_spacing(p)
            p.paragraph_format.left_indent = Pt(24)
            text = " ".join(ln.lstrip("> ") for ln in raw_block.splitlines())
            add_runs(p, inline_runs(text), color=NAVY, force_italic=True)
            continue
        block = " ".join(raw_block.splitlines()).strip()
        p = doc.add_paragraph()
        style_spacing(p)
        if block.startswith("## "):
            seen_q = True
            qm = re.match(r"## (Q\d+):\s*(.*)", block)
            if qm:
                r = p.add_run(qm.group(1) + "  ")
                r.bold = True
                r.font.size = Pt(12)
                r.font.color.rgb = NAVY
                r2 = p.add_run(smarten(qm.group(2)))
                r2.bold = True
                r2.font.size = Pt(12)
            else:
                r = p.add_run(smarten(block[3:]))
                r.bold = True
                r.font.size = Pt(12)
                r.font.color.rgb = NAVY
        elif block.startswith('<p class="source">'):
            text = re.sub(r"</?p[^>]*>", "", block)
            label, rest = text.split(": ", 1)
            r = p.add_run(label + ": ")
            r.bold, r.italic = True, True
            r.font.size = Pt(9)
            r.font.color.rgb = GRAY
            add_runs(p, inline_runs(rest), size=Pt(9), color=GRAY, force_italic=True)
        elif block.startswith('<p class="paraphrase">'):
            text = re.sub(r"</?p[^>]*>|</?strong>", "", block)
            label, rest = text.split(": ", 1)
            r = p.add_run(label + ": ")
            r.bold, r.italic = True, True
            r.font.color.rgb = NAVY
            add_runs(p, inline_runs(rest), color=NAVY, force_italic=True)
        elif not seen_q and block.startswith("*") and block.endswith("*"):
            add_runs(p, inline_runs(block[1:-1]), size=Pt(12), color=GRAY,
                     force_italic=True)
        else:
            add_runs(p, inline_runs(block))

    suffix = "FINAL" if "--final" in sys.argv else "DRAFT"
    safe_title = title.replace("?", "").replace(":", " —").strip()
    out = Path.home() / "Desktop" / f"{safe_title} — {suffix}.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
